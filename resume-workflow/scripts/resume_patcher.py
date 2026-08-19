#!/usr/bin/env python3
"""
Resume Patcher v4 — 简历定制脚本
新增 delete_bullet 操作：保留经历块的前 N 个 bullet，删除后面的。

用法:
    python resume_patcher.py --input 简历.docx --instructions instructions.json --output 新简历.docx
"""

import json
import argparse
from docx import Document
from docx.oxml.ns import qn


def is_heading_paragraph(para):
    """判断段落是否为标题/加粗段落（块的开头）"""
    if para.style.name.startswith("Heading"):
        return True
    for run in para.runs:
        if run.bold:
            return True
    for r in para._element.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None and rPr.find(qn("w:b")) is not None:
            return True
    return False


SECTION_KEYWORDS = [
    "教育经历", "项目经历", "工作经历", "基本技能", "其他",
    "Educational Experience", "Education",
    "Research Experience", "Industry Experience",
    "Basic Skills", "Others", "Campus Experience",
    "校园经历", "技能", "Skills"
]


def is_section_title(para):
    """判断段落是否为章节标题"""
    text = para.text.strip()
    if not is_heading_paragraph(para):
        return False
    for kw in SECTION_KEYWORDS:
        if text == kw or text.startswith(kw):
            return True
    return False


def find_section_paragraphs(doc, section_name):
    """找到章节内的所有段落"""
    in_section = False
    section_paras = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.startswith(section_name) and is_section_title(para):
            in_section = True
            continue

        if in_section and is_section_title(para):
            in_section = False

        if in_section:
            section_paras.append(para)

    return section_paras


def group_blocks(paragraphs):
    """将段落列表分组为块"""
    blocks = []
    current_block = None

    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if is_heading_paragraph(para):
            if current_block is not None:
                blocks.append(current_block)
            current_block = {
                "headers": [para],
                "bullets": [],
            }
        else:
            if current_block is None:
                continue
            current_block["bullets"].append(para)

    if current_block is not None:
        blocks.append(current_block)

    return blocks


def block_header_text(block):
    return " ".join([h.text for h in block["headers"]])


def block_matches(block, identifier):
    header = block_header_text(block)
    return identifier in header


def move_paragraph_after(para, target_para):
    parent = para._element.getparent()
    target_parent = target_para._element.getparent()
    if parent is not target_parent:
        return

    try:
        current_idx = list(parent).index(para._element)
        target_idx = list(parent).index(target_para._element)
        if current_idx == target_idx + 1:
            return
    except ValueError:
        pass

    parent.remove(para._element)
    target_idx = list(parent).index(target_para._element)
    parent.insert(target_idx + 1, para._element)


def remove_paragraph(para):
    para._element.getparent().remove(para._element)


def rewrite_paragraph(para, old_text, new_text):
    full_text = para.text
    if old_text not in full_text:
        return False

    new_full_text = full_text.replace(old_text, new_text)

    if not para.runs:
        para.add_run(new_full_text)
        return True

    first_run = para.runs[0]
    saved_format = {
        "name": first_run.font.name,
        "size": first_run.font.size,
        "bold": first_run.bold,
        "italic": first_run.italic,
        "underline": first_run.underline,
        "color": first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None,
    }

    for run in list(para.runs):
        run._element.getparent().remove(run._element)

    new_run = para.add_run(new_full_text)
    if saved_format["name"]:
        new_run.font.name = saved_format["name"]
    if saved_format["size"]:
        new_run.font.size = saved_format["size"]
    if saved_format["bold"] is not None:
        new_run.bold = saved_format["bold"]
    if saved_format["italic"] is not None:
        new_run.italic = saved_format["italic"]
    if saved_format["underline"] is not None:
        new_run.underline = saved_format["underline"]
    if saved_format["color"]:
        new_run.font.color.rgb = saved_format["color"]

    return True


def execute_reorder(doc, section_name, items):
    section_title = None
    for para in doc.paragraphs:
        if para.text.strip().startswith(section_name) and is_section_title(para):
            section_title = para
            break

    if section_title is None:
        print(f"  [警告] 未找到章节标题: {section_name}")
        return

    section_paras = find_section_paragraphs(doc, section_name)
    blocks = group_blocks(section_paras)

    if not blocks:
        print(f"  [警告] 章节内无内容: {section_name}")
        return

    print(f"  [信息] 章节 '{section_name}' 内发现 {len(blocks)} 个块")

    ordered_blocks = []
    remaining = list(blocks)

    for item in items:
        ident = item["identifier"]
        found = False
        for i, block in enumerate(remaining):
            if block_matches(block, ident):
                ordered_blocks.append(block)
                remaining.pop(i)
                found = True
                break
        if not found:
            print(f"  [警告] 未找到块: {ident}")

    ordered_blocks.extend(remaining)

    current_anchor = section_title
    for block in ordered_blocks:
        all_paras = block["headers"] + block["bullets"]
        for para in all_paras:
            move_paragraph_after(para, current_anchor)
            current_anchor = para

    print(f"  [完成] 重排完成，指定 {len(items)} 个 + 剩余 {len(remaining)} 个")


def execute_delete(doc, section_name, items):
    if section_name:
        section_paras = find_section_paragraphs(doc, section_name)
        blocks = group_blocks(section_paras)
    else:
        blocks = group_blocks(doc.paragraphs)

    identifiers = [item["identifier"] for item in items]
    removed = 0

    for block in blocks:
        if any(block_matches(block, ident) for ident in identifiers):
            for para in block["headers"] + block["bullets"]:
                remove_paragraph(para)
            removed += 1

    print(f"  [完成] 删除 {removed} 个块")


def execute_rewrite(doc, section_name, identifier, changes):
    section_paras = find_section_paragraphs(doc, section_name)
    blocks = group_blocks(section_paras)

    rewritten = 0
    for block in blocks:
        if block_matches(block, identifier):
            for para in block["bullets"]:
                for change in changes:
                    if rewrite_paragraph(para, change["original"], change["replacement"]):
                        rewritten += 1
                        break
            break

    print(f"  [完成] 重写 {rewritten} 个段落")


def execute_delete_bullet(doc, section_name, identifier, keep_first):
    """新增：保留经历块的前 N 个 bullet，删除后面的"""
    section_paras = find_section_paragraphs(doc, section_name)
    blocks = group_blocks(section_paras)

    for block in blocks:
        if block_matches(block, identifier):
            bullets = block["bullets"]
            if len(bullets) <= keep_first:
                print(f"  [信息] '{identifier}' 只有 {len(bullets)} 个 bullet，无需删除")
                return

            # 删除第 keep_first 个之后的所有 bullet
            to_delete = bullets[keep_first:]
            for para in to_delete:
                remove_paragraph(para)

            print(f"  [完成] 删除 '{identifier}' 的 {len(to_delete)} 个 bullet，保留前 {keep_first} 个")
            return

    print(f"  [警告] 未找到块: {identifier}")


EXECUTION_ORDER = ["reorder", "delete", "delete_bullet", "rewrite"]


def apply_instructions(doc, instructions):
    """应用所有指令。执行顺序固定为 reorder -> delete -> delete_bullet -> rewrite，
    与 JSON 中 operations 数组的书写顺序无关（同类型内部保持原有相对顺序），
    以确保 rewrite 不会作用于已被 delete/delete_bullet 移除的段落。"""
    operations = instructions.get("operations", [])
    ordered_ops = sorted(
        operations,
        key=lambda op: EXECUTION_ORDER.index(op["type"]) if op["type"] in EXECUTION_ORDER else len(EXECUTION_ORDER)
    )

    for op in ordered_ops:
        op_type = op["type"]
        section = op.get("section")

        if op_type == "reorder":
            print(f"[操作] 重排: {section}")
            execute_reorder(doc, section, op["items"])

        elif op_type == "delete":
            scope = section if section else "全局"
            print(f"[操作] 删除: {scope}")
            execute_delete(doc, section, op["items"])

        elif op_type == "rewrite":
            print(f"[操作] 重写: {section} -> {op['identifier']}")
            execute_rewrite(doc, section, op["identifier"], op["changes"])

        elif op_type == "delete_bullet":
            print(f"[操作] 删除多余 bullet: {section} -> {op['identifier']} (保留前 {op.get('keep_first', 3)} 个)")
            execute_delete_bullet(doc, section, op["identifier"], op.get("keep_first", 3))


def main():
    parser = argparse.ArgumentParser(description="Resume Patcher — 简历定制脚本")
    parser.add_argument("--input", required=True, help="原始简历 docx 路径")
    parser.add_argument("--instructions", required=True, help="JSON 指令文件路径")
    parser.add_argument("--output", required=True, help="输出简历 docx 路径")
    args = parser.parse_args()

    print(f"读取原始文件: {args.input}")
    doc = Document(args.input)

    print(f"读取指令文件: {args.instructions}")
    with open(args.instructions, "r", encoding="utf-8") as f:
        instructions = json.load(f)

    print("\n开始应用指令...")
    apply_instructions(doc, instructions)

    print(f"\n保存到: {args.output}")
    doc.save(args.output)
    print("完成！")


if __name__ == "__main__":
    main()
